"""
Codebase to Text Converter

This module provides functionality to convert codebases (folder structures with files)
into a single text file, markdown file (.md), or Microsoft Word document (.docx), while preserving folder
structure and file contents. It supports advanced exclusion patterns and GitHub repositories.
"""

import os
import argparse
import shutil
import fnmatch
import tempfile
import sys
import base64
import io
import logging
import uuid
import abc
from pathlib import Path
from xml.sax.saxutils import quoteattr
from typing import List, Optional, Set, Tuple, Any


from codebase_convert.utils import clone_github_repo, process_files_with_strategy, is_image_file, compress_image

class OutputFormatter(abc.ABC):
    @abc.abstractmethod
    def format_file_success(self, file_path: str, base_path: str, file_content: str, ext: str, ai_optimize: bool) -> str:
        pass
        
    @abc.abstractmethod
    def format_file_error(self, file_path: str, base_path: str, error: Exception) -> str:
        pass
        
    @abc.abstractmethod
    def combine_text(self, folder_structure: str, file_contents: str) -> str:
        pass

    def save_file(self, text: str, output_path: str, verbose: bool) -> None:
        with open(output_path, "w", encoding="utf-8") as file:
            file.write(text)

class TxtFormatter(OutputFormatter):
    def format_file_success(self, file_path: str, base_path: str, file_content: str, ext: str, ai_optimize: bool) -> str:
        rel_path = os.path.relpath(file_path, base_path)
        if ai_optimize:
            content = f"<file path=\"{rel_path}\">\n"
            content += f"{file_content}\n"
            content += f"\n</file>\n"
            return content

        content = f"\n\n{rel_path}\n"
        content += f"File type: {ext or 'no extension'}\n"
        content += f"{file_content}"
        content += f"\n\n{'-' * 50}\nFile End\n{'-' * 50}\n"
        return content
        
    def format_file_error(self, file_path: str, base_path: str, error: Exception) -> str:
        rel_path = os.path.relpath(file_path, base_path)
        content = f"\n\n{rel_path}\n"
        content += f"File type: {os.path.splitext(file_path)[1] or 'no extension'}\n"
        content += f"[Error: Could not process file - {str(error)}]"
        content += f"\n\n{'-' * 50}\nFile End\n{'-' * 50}\n"
        return content
        
    def combine_text(self, folder_structure: str, file_contents: str) -> str:
        folder_structure_header = "Folder Structure"
        file_contents_header = "File Contents"
        delimiter = "-" * 50
        return (f"{folder_structure_header}\n{delimiter}\n{folder_structure}\n\n"
                f"{file_contents_header}\n{delimiter}\n{file_contents}")

class DocxFormatter(TxtFormatter):
    def __init__(self):
        self.marker_start = f"(IMAGE_MARKER_{uuid.uuid4().hex})"
        self.marker_end = f"(/IMAGE_MARKER_{uuid.uuid4().hex})"

    def save_file(self, text: str, output_path: str, verbose: bool) -> None:
        doc = Document()
        segments = text.split(self.marker_start)
        
        if segments[0].strip():
            doc.add_paragraph(segments[0])
            
        for segment in segments[1:]:
            if self.marker_end in segment:
                img_path, text_content = segment.split(self.marker_end, 1)
                img_path = str(img_path.strip())
                if verbose:
                    logger.debug(f"Loading image from: {img_path}")
                if not os.path.exists(img_path):
                    if verbose:
                        logger.warning(f"Image file not found at: {img_path}")
                    doc.add_paragraph(f"[Missing image: {img_path}]")
                else:
                    try:
                        doc.add_picture(img_path, width=Inches(6))
                    except Exception as e:
                        if verbose:
                            logger.error(f"Error adding image {img_path}: {e}")
                        doc.add_paragraph(f"[Error: Could not add image - {str(e)}]")
                
                if text_content.strip():
                    doc.add_paragraph(text_content)
        doc.save(output_path)

class MdFormatter(OutputFormatter):
    def _get_lang_from_ext(self, ext: str) -> str:
        extension_map = {
            '.py': 'python', '.js': 'javascript', '.ts': 'typescript', 
            '.jsx': 'jsx', '.tsx': 'tsx', '.html': 'html', '.css': 'css', 
            '.json': 'json', '.md': 'markdown', '.yml': 'yaml', '.yaml': 'yaml',
            '.sh': 'bash', '.go': 'go', '.rs': 'go', '.java': 'java', 
            '.cpp': 'java', '.c': 'c', '.h': 'c', '.cs': 'csharp', 
            '.rb': 'ruby', '.php': 'php', '.sql': 'sql', '.xml': 'xml'
        }
        return extension_map.get(ext.lower(), '')

    def format_file_success(self, file_path: str, base_path: str, file_content: str, ext: str, ai_optimize: bool) -> str:
        rel_path = os.path.relpath(file_path, base_path)
        lang = self._get_lang_from_ext(ext)
        return f"\n### File: `{rel_path}`\n\n```{lang}\n{file_content}\n```\n"

    def format_file_error(self, file_path: str, base_path: str, error: Exception) -> str:
        rel_path = os.path.relpath(file_path, base_path)
        return f"\n### File: `{rel_path}`\n\n```text\n[Error: Could not process file - {str(error)}]\n```\n"
        
    def combine_text(self, folder_structure: str, file_contents: str) -> str:
        folder_structure_header = "Folder Structure"
        file_contents_header = "File Contents"
        return (f"# {folder_structure_header}\n\n```text\n{folder_structure}```\n\n"
                f"# {file_contents_header}\n{file_contents}")

def get_formatter(output_type: str) -> OutputFormatter:
    if output_type == 'md':
        return MdFormatter()
    elif output_type == 'docx':
        return DocxFormatter()
    elif output_type == 'txt':
        return TxtFormatter()
    else:
        raise ValueError(f"Invalid output type: {output_type}. Supported types: txt, docx, md")

import git
import pathspec
from docx import Document
from docx.shared import Inches

try:
    import tiktoken
    HAS_TIKTOKEN = True
except ImportError:
    HAS_TIKTOKEN = False

logger = logging.getLogger("codebase_convert")

MAX_TEXT_FILE_BYTES = int(os.environ.get("CODEBASE_CONVERT_MAX_TEXT_FILE_BYTES", "1000000"))
BINARY_SAMPLE_BYTES = 8192

class CodebaseConvert:
    """
    Convert codebase to text with advanced exclusion patterns.

    This class provides functionality to convert local directories or GitHub repositories
    into text or DOCX files while supporting sophisticated exclusion patterns including
    wildcards, directory patterns, and .exclude file support.
    """

    def __init__(self, input_path: str, output_path: str, output_type: str, verbose: bool = False,
                 exclude_hidden: bool = False, exclude: Optional[List[str]] = None, 
                 ai_optimize: bool = True, strip_comments: bool = False):
        """
        Initialize CodebaseConvert converter.
        
        Args:
            input_path: Path to local directory or GitHub repository URL
            output_path: Path for the output file
            output_type: Output format ('txt', 'md' or 'docx')
            verbose: Enable detailed logging (default: False)
            exclude_hidden: Exclude hidden files and directories (default: False)
            exclude: List of exclusion patterns (default: None)
            ai_optimize: Optimize output for AI readability and compress size (default: True)
            strip_comments: Remove lines with comment prefixes (default: False)
        """
        self.input_path = input_path
        self.output_path = output_path
        self.output_type = output_type
        self.config = {
            'verbose': verbose,
            'exclude_hidden': exclude_hidden,
            'ai_optimize': ai_optimize,
            'strip_comments': strip_comments,
        }
        
        # Configure logging
        if verbose:
            logging.basicConfig(level=logging.DEBUG, format='%(levelname)s: %(message)s')
        else:
            logging.basicConfig(level=logging.INFO, format='%(message)s')

        self.formatter = get_formatter(self.output_type)

        # Initialize exclusion patterns
        self.exclude_patterns: Set[str] = set()
        self.excluded_files_count = 0
        self.gitignore_spec: Optional[pathspec.PathSpec] = None

        # Load exclusion patterns from various sources
        self._load_exclusion_patterns(exclude)

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        pass

    @property
    def verbose(self) -> bool:
        """Get verbose setting"""
        return self.config['verbose']

    @property
    def exclude_hidden(self) -> bool:
        """Get exclude_hidden setting"""
        return self.config['exclude_hidden']

    def _load_exclusion_patterns(self, exclude_args: Optional[List[str]]):
        """Load exclusion patterns from CLI args, defaults and .exclude file."""
        self._add_cli_patterns(exclude_args)
        self._add_default_patterns()
        self._add_file_patterns()
        self._load_gitignore()

        if self.verbose and self.exclude_patterns:
            logger.debug(f"Active exclusion patterns: {sorted(self.exclude_patterns)}")

    def _load_gitignore(self, base_path: Optional[str] = None):
        """Load patterns from local .gitignore file if present."""
        search_path = base_path if base_path else (self.input_path if not self.is_github_repo() else '.')
        gitignore_path = os.path.join(search_path, '.gitignore')
        
        if not os.path.exists(gitignore_path):
            return

        try:
            with open(gitignore_path, "r", encoding="utf-8") as f:
                self.gitignore_spec = pathspec.PathSpec.from_lines("gitignore", f)

            if self.verbose:
                logger.debug(f"Loaded gitignore patterns from {gitignore_path}")

        except Exception as e:
            if self.verbose:
                logger.warning(f"Could not read .gitignore file: {e}")

    def _add_cli_patterns(self, exclude_args: Optional[List[str]]):
        """Add patterns provided via command line."""
        if not exclude_args:
            return
        for pattern in exclude_args:
            for p in pattern.split(','):
                p = p.strip()
                if p:
                    self.exclude_patterns.add(p)

    
    def _add_default_patterns(self):
        """Add default exclusion patterns for common files/folders."""
        default_excludes = {
            '.git/', '.git/**', '.github/', '.gitlab/', '.hg/', '.hg/**',
            '__pycache__/', '**/__pycache__/**',
            '*.pyc', '*.pyo', '*.pyd',
            '.venv/', 'venv/', 'env/', '.env',
            '.env.local',
            'node_modules/',
            '.DS_Store',
            '*.log', '*.tmp',
            '.pytest_cache/',
            '.coverage',
            'build/', 'dist/',
            '*.egg-info/',
            '.vscode/', '*-lock.json', '*-lock.yaml',
            '.mypy_cache/', 'server_uploads/', '*.jsonl', '.env.keys',
            '.netlify/', '.vercel/', '.aws-sam/', '.serverless/', '.azure/', '.gcloud/',
            'coverage/', 'logs/', 'debug.log', 'debug.log.*',
            '.idea/', '*.iml', '*.iws', '*.ipr',
            '.vs/', '*.sln', '*.suo', '*.vcxproj*',
            '.gradle/', 'build/', 'out/',
            '.vscode-test/',
        }
        if self.config.get('ai_optimize', False):
            media_excludes = {
                '*.mp3', '*.mp4', '*.wav', '*.avi', '*.mov', '*.flv', '*.wmv', '*.webm', '*.ogg',
                '*.pdf', '*.eot', '*.ttf', '*.woff', '*.woff2'
            }
            default_excludes.update(media_excludes)
            
        self.exclude_patterns.update(default_excludes)
    def _add_file_patterns(self, base_path: Optional[str] = None):
        """Load patterns from a .exclude file if present."""
        search_path = base_path if base_path else (self.input_path if not self.is_github_repo() else '.')
        exclude_file_path = os.path.join(search_path, '.exclude')
        
        if not os.path.exists(exclude_file_path):
            return
        try:
            with open(exclude_file_path, 'r', encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if line and not line.startswith('#'):
                        self.exclude_patterns.add(line)
            if self.verbose:
                logger.debug(f"Loaded exclusion patterns from {exclude_file_path}")
        except (OSError, UnicodeDecodeError) as e:
            if self.verbose:
                logger.warning(f"Warning: Could not read .exclude file: {e}")

    def _normalize_path(self, file_path: str, base_path: str) -> str:
        """Normalize path for pattern matching"""
        try:
            # Get relative path from base
            rel_path = os.path.relpath(file_path, base_path)
            # Convert to forward slashes for consistent pattern matching
            return rel_path.replace(os.sep, '/')
        except ValueError:
            # If relative path can't be computed, use absolute path
            return file_path.replace(os.sep, '/')

    def _should_exclude(self, file_path: str, base_path: str) -> bool:
        """Check if file/directory should be excluded based on patterns"""
        # Handle hidden files if exclude_hidden is True
        if self.exclude_hidden and self._is_hidden_file(file_path):
            return True

        # Normalize the path for pattern matching
        normalized_path = self._normalize_path(file_path, base_path)
        filename = os.path.basename(file_path)

        # Check against gitignore if present
        if self.gitignore_spec and self.gitignore_spec.match_file(normalized_path):
            return True

        if not self.exclude_patterns:
            return False

        # Check against all exclusion patterns
        for pattern in self.exclude_patterns:
            if self._pattern_matches(pattern, normalized_path, filename):
                return True

        return False

    def _pattern_matches(self, pattern: str, normalized_path: str, filename: str) -> bool:
        """Return True if the pattern matches the given file."""
        if fnmatch.fnmatch(filename, pattern):
            return True
        if fnmatch.fnmatch(normalized_path, pattern):
            return True
        if pattern.endswith('/'):
            if self._dir_pattern_match(pattern.rstrip('/'), normalized_path):
                return True
        if '**' in pattern:
            if self._recursive_pattern_match(pattern, normalized_path):
                return True

        return False

    def _dir_pattern_match(self, dir_pattern: str, normalized_path: str) -> bool:
        """Check directory style pattern."""
        for part in normalized_path.split('/'):
            if fnmatch.fnmatch(part, dir_pattern):
                return True
        return False

    def _recursive_pattern_match(self, pattern: str, normalized_path: str) -> bool:
        """Handle patterns that include ** for recursion."""
        recursive_pattern = pattern.replace('**/', '').replace('**', '*')
        if fnmatch.fnmatch(normalized_path, recursive_pattern):
            return True
        path_parts = normalized_path.split('/')
        for i in range(len(path_parts)):
            partial_path = '/'.join(path_parts[i:])
            if fnmatch.fnmatch(partial_path, recursive_pattern):
                return True
        return False

    def _parse_folder(self, folder_path: str) -> str:
        """Parse folder structure, respecting exclusion patterns"""
        tree = ""
        excluded_dirs: Set[str] = set()

        for root, dirs, files in os.walk(folder_path):
            if self._handle_excluded_directory(root, folder_path, excluded_dirs):
                continue

            if self._skip_excluded_dir(root, excluded_dirs):
                continue

            self._filter_excluded_directories(dirs, root, folder_path)
            tree += self._generate_directory_entry(root, folder_path)
            tree += self._generate_file_entries(files, root, folder_path)

        self._log_tree_results(tree)
        return tree

    def _handle_excluded_directory(self, root: str, folder_path: str, excluded_dirs: Set[str]) -> bool:
        """Handle directory exclusion and return True if directory should be skipped"""
        if self._should_exclude(root, folder_path):
            if self.verbose:
                logger.debug(f"Excluding directory: {root}")
            self.excluded_files_count += 1
            excluded_dirs.add(root)
            return True
        return False

    def _filter_excluded_directories(self, dirs: List[str], root: str, folder_path: str):
        """Filter out excluded directories from the dirs list"""
        original_dirs = dirs[:]
        dirs[:] = []
        for d in original_dirs:
            dir_path = os.path.join(root, d)
            if not self._should_exclude(dir_path, folder_path):
                dirs.append(d)
            elif self.verbose:
                logger.debug(f"Excluding directory from tree: {dir_path}")
                self.excluded_files_count += 1

        # Apply hidden file exclusion
        if self.exclude_hidden:
            dirs[:] = [d for d in dirs if not self._is_hidden_file(os.path.join(root, d))]
            
    def _generate_directory_entry(self, root: str, folder_path: str) -> str:
        """Generate tree entry for a directory"""
        level = root.replace(folder_path, '').count(os.sep)
        indent = ' ' * 4 * level
        return f'{indent}{os.path.basename(root)}/\n'

    def _generate_file_entries(self, files: List[str], root: str, folder_path: str) -> str:
        """Generate tree entries for files in a directory"""
        tree = ""
        level = root.replace(folder_path, '').count(os.sep)
        subindent = ' ' * 4 * (level + 1)
        for f in files:
            file_path = os.path.join(root, f)
            if not self._should_exclude(file_path, folder_path):
                tree += f'{subindent}{f}\n'
            elif self.verbose:
                logger.debug(f"Excluding file from tree: {file_path}")
                self.excluded_files_count += 1
        return tree

    def _log_tree_results(self, tree: str):
        """Log tree generation results if verbose mode is enabled"""
        if self.verbose:
            logger.debug(f"The file tree to be processed:\n{tree}")
            logger.debug(f"Total excluded items: {self.excluded_files_count}")

    def _is_within_base(self, file_path: str, base_path: str) -> bool:
        try:
            resolved_file = Path(file_path).resolve(strict=True)
            resolved_base = Path(base_path).resolve(strict=True)
            resolved_file.relative_to(resolved_base)
            return True
        except (OSError, ValueError):
            return False


    def _is_probably_binary(self, file_path: str) -> bool:
        try:
            with open(file_path, "rb") as file:
                sample = file.read(BINARY_SAMPLE_BYTES)

            return b"\0" in sample

        except OSError:
            return True

    def _get_file_contents(self, file_path: str) -> str:
        """Read file contents with size limits and binary detection."""
        try:
            file_size = os.path.getsize(file_path)
        except OSError as e:
            return f"[Error reading file metadata: {str(e)}]"

        if file_size > MAX_TEXT_FILE_BYTES:
            return f"[Skipped: file exceeds {MAX_TEXT_FILE_BYTES:,} bytes]"

        if self._is_probably_binary(file_path):
            return "[Skipped: binary file]"

        try:
            with open(file_path, "r", encoding="utf-8") as file:
                return file.read()

        except UnicodeDecodeError:
            try:
                with open(file_path, "r", encoding="utf-8", errors="replace") as file:
                    return file.read()
            except OSError as e:
                return f"[Error reading file: {str(e)}]"

        except OSError as e:
            return f"[Error reading file: {str(e)}]"

    def _is_hidden_file(self, file_path: str) -> bool:
        """Check if file/directory is hidden"""
        components = os.path.normpath(file_path).split(os.sep)
        for c in components:
            if c.startswith((".", "__")):
                return True
        return False

    def _skip_excluded_dir(self, root: str, excluded_dirs: Set[str]) -> bool:
        """Return True if the directory is inside an excluded path."""
        for excluded_dir in excluded_dirs:
            if root.startswith(excluded_dir):
                return True
        return False

    def _process_files(self, path: str) -> str:
        """Process files based on local or remote strategy, respecting exclusion patterns"""
        content_pieces, processed_count = process_files_with_strategy(
            self.is_github_repo(),
            path,
            self._should_exclude,
            self._handle_directory_exclusion,
            self._filter_directories_for_processing,
            self._process_single_file
        )

        if self.verbose:
            logger.info(f"Processed {processed_count} files")

        return "".join(content_pieces)

    def _handle_directory_exclusion(self, root: str, path: str, excluded_dirs: Set[str]) -> bool:
        """Handle directory exclusion for file processing"""
        if self._should_exclude(root, path):
            if self.verbose:
                logger.debug(f"Skipping excluded directory: {root}")
            excluded_dirs.add(root)
            return True
        return False

    def _filter_directories_for_processing(self, dirs: List[str], root: str, path: str):
        """Filter directories to avoid processing excluded ones"""
        original_dirs = dirs[:]
        dirs[:] = []
        for d in original_dirs:
            dir_path = os.path.join(root, d)
            if not self._should_exclude(dir_path, path):
                dirs.append(d)

    def _process_single_file(self, file: str, root: str, path: str) -> Optional[str]:
        """Process a single file and return its content or None if excluded."""
        file_path = os.path.join(root, file)

        if not self._is_within_base(file_path, path):
            if self.verbose:
                logger.warning(f"Skipping file outside base path: {file_path}")
            return None

        if os.path.islink(file_path):
            if self.verbose:
                logger.warning(f"Skipping symlinked file: {file_path}")
            return None

        if self._should_exclude(file_path, path):
            if self.verbose:
                logger.debug(f"Skipping excluded file: {file_path}")
            return None

        if self.verbose:
            logger.debug(f"Processing: {file_path}")

        try:
            if self.output_type == "docx" and is_image_file(file_path):
                return f"\n\n{self.formatter.marker_start}{os.path.abspath(file_path)}{self.formatter.marker_end}\n"

            if self.config.get("ai_optimize", False) and is_image_file(file_path):
                rel_path = os.path.relpath(file_path, path)

                blob_bytes, mime_type = compress_image(file_path, verbose=self.verbose)

                if not blob_bytes or not mime_type:
                    return f"[Skipped image: could not safely process {rel_path}]\n"

                blob = base64.b64encode(blob_bytes).decode("utf-8")

                return (
                    f'<image path={quoteattr(rel_path)} '
                    f'type={quoteattr(mime_type)} '
                    f'compressed="true">\n'
                    f"{blob}\n"
                    f"</image>\n\n"
                )

            return self._format_file_content(file_path, path)

        except (OSError, UnicodeDecodeError) as e:
            return self._format_file_error(file_path, path, e)

    def _optimize_for_ai(self, content: str) -> str:
        lines = content.splitlines()
        cleaned = []
        comment_prefixes = ('//', '#', '/*', '*')
        strip_comments = self.config.get('strip_comments', False)
        
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue
            if strip_comments and stripped.startswith(comment_prefixes):
                continue
            cleaned.append(line)
        return '\n'.join(cleaned)

    def _get_lang_from_ext(self, ext: str) -> str:
        extension_map = {
            '.py': 'python', '.js': 'javascript', '.ts': 'typescript', 
            '.jsx': 'jsx', '.tsx': 'tsx', '.html': 'html', '.css': 'css', 
            '.json': 'json', '.md': 'markdown', '.yml': 'yaml', '.yaml': 'yaml',
            '.sh': 'bash', '.go': 'go', '.rs': 'go', '.java': 'java', 
            '.cpp': 'java', '.c': 'c', '.h': 'c', '.cs': 'csharp', 
            '.rb': 'ruby', '.php': 'php', '.sql': 'sql', '.xml': 'xml'
        }
        return extension_map.get(ext.lower(), '')

    def _format_file_content(self, file_path: str, base_path: str) -> str:
        """Format successful file content for output"""
        file_content = self._get_file_contents(file_path)
        ext = os.path.splitext(file_path)[1]

        if self.config.get('ai_optimize', True):
            file_content = self._optimize_for_ai(file_content)
        
        return self.formatter.format_file_success(file_path, base_path, file_content, ext, self.config.get('ai_optimize', True))

    def _format_file_error(self, file_path: str, base_path: str, error: Exception) -> str:
        """Format file processing error for output"""
        if self.verbose:
            logger.error(f"Error processing {file_path}: {error}")
            
        return self.formatter.format_file_error(file_path, base_path, error)

    def get_text(self) -> str:
        """Generate the combined text output."""
        if self.is_github_repo():
            temp_folder_path = clone_github_repo(self.input_path, self.verbose)

            try:
                self._add_file_patterns(temp_folder_path)
                self._load_gitignore(temp_folder_path)

                folder_structure = self._parse_folder(temp_folder_path)
                file_contents = self._process_files(temp_folder_path)

            finally:
                shutil.rmtree(temp_folder_path, ignore_errors=True)

        else:
            folder_structure = self._parse_folder(self.input_path)
            file_contents = self._process_files(self.input_path)

        return self.formatter.combine_text(folder_structure, file_contents)

        # Section headers
        return self.formatter.combine_text(folder_structure, file_contents)

    # Estimate token method removed. Use codebase_convert.utils.estimate_tokens instead.

    def get_file(self, text_output: Optional[str] = None):
        """Generate and save the output file"""
        if self.verbose:
            logger.info("Generating text layout...")
        text = text_output if text_output is not None else self.get_text()

        self.formatter.save_file(text, self.output_path, self.verbose)

        if self.verbose:
            logger.info(f"Output saved to: {self.output_path}")
            
        from codebase_convert.utils import estimate_tokens
        token_count = estimate_tokens(text, verbose=self.verbose)
        if token_count is not None:
            logger.info(f"Estimated token count (cl100k_base): ~{token_count:,}")

    #### GitHub Support ####

    def is_github_repo(self) -> bool:
        """Check if input path is a GitHub repository URL"""
        return (self.input_path.startswith("https://github.com/") or
                self.input_path.startswith("git@github.com:"))



def main():
    """Main CLI entry point"""
    parser = argparse.ArgumentParser(
        description="Convert codebase to text with exclusion support.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s --input ./my_project --output output.txt --output_type txt
  %(prog)s --input https://github.com/user/repo --output repo.md --output_type md
  %(prog)s --input ./project --output out.txt --output_type txt --exclude "*.log,temp/,**/__pycache__/**"
        """
    )

    parser.add_argument("--input", help="Input path (folder or GitHub URL)", required=True)
    parser.add_argument("--output", help="Output file path", required=True)
    parser.add_argument("--output_type", help="Output file type (txt, md, or docx)",
                       choices=["txt", "docx", "md"], default="txt")
    parser.add_argument("--exclude", help="Exclude patterns (can be used multiple times)",
                       action="append", default=[])
    parser.add_argument("--exclude_hidden", help="Exclude hidden files and folders",
                       action="store_true")
    parser.add_argument("--verbose", help="Show detailed processing information",
                       action="store_true")
    parser.add_argument("--no_ai_optimize", help="Disable AI optimization (formatting, removing empty lines)",
                       action="store_true")
    parser.add_argument("--strip_comments", help="Remove comment lines from code (used with ai_optimize)",
                       action="store_true")

    args = parser.parse_args()

    try:
        with CodebaseConvert(
            input_path=args.input,
            output_path=args.output,
            output_type=args.output_type,
            verbose=args.verbose,
            exclude_hidden=args.exclude_hidden,
            exclude=args.exclude,
            ai_optimize=not args.no_ai_optimize,
            strip_comments=args.strip_comments
        ) as code_to_text:

            code_to_text.get_file()

        logger.info("✅ Conversion completed successfully!")

    except (OSError, ValueError, git.GitCommandError) as e:
        logger.error(f"❌ Error: {e}")
        return 1

    return 0


if __name__ == "__main__":
    sys.exit(main())

